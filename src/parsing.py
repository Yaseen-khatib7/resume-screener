from pathlib import Path
import io
import os
import re
import zipfile
import fitz  # PyMuPDF
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

try:
    from PIL import Image, ImageFilter, ImageOps
    import pytesseract
except ImportError:  # pragma: no cover - optional OCR dependency
    Image = None
    ImageFilter = None
    ImageOps = None
    pytesseract = None


CHAR_REPLACEMENTS = {
    "\u00a0": " ",
    "\u200b": "",
    "\ufeff": "",
    "â€¢": "•",
    "â€“": "-",
    "â€”": "-",
    "â€": "\"",
    "’": "'",
    "“": "\"",
    "”": "\"",
}
PDF_OCR_FALLBACK_MIN_CHARS = 180
WINDOWS_TESSERACT_PATHS = (
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
)


def _normalize_extracted_text(text: str) -> str:
    normalized = text or ""
    for source, target in CHAR_REPLACEMENTS.items():
        normalized = normalized.replace(source, target)
    normalized = normalized.replace("\r\n", "\n").replace("\r", "\n")
    # Repair fragmented email tokens without crossing line boundaries, otherwise
    # headings like "Skills" can be glued onto the email domain.
    normalized = re.sub(r"([A-Za-z0-9._%+-]+)[ \t]*@[ \t]*([A-Za-z0-9.-]+)", r"\1@\2", normalized)
    normalized = re.sub(r"([A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+)[ \t]*\.[ \t]*([A-Za-z]{2,4})\b", r"\1.\2", normalized)
    normalized = re.sub(r"[ \t\f\v]+", " ", normalized)
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _format_links(links: list[str]) -> str:
    cleaned = []
    seen = set()
    for link in links:
        value = (link or "").strip()
        if not value:
            continue
        key = value.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(value)
    if not cleaned:
        return ""
    return "\n".join(f"HYPERLINK: {link}" for link in cleaned)


def _dedupe_lines(lines: list[str]) -> list[str]:
    cleaned: list[str] = []
    seen = set()
    for line in lines:
        value = re.sub(r"\s+", " ", (line or "").strip())
        if not value:
            continue
        key = value.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(value)
    return cleaned


def _read_docx_xml_text(data: bytes) -> tuple[list[str], list[str]]:
    text_parts: list[str] = []
    links: list[str] = []

    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        names = archive.namelist()

        xml_targets = [
            name
            for name in names
            if name.startswith("word/")
            and name.endswith(".xml")
            and not name.endswith(".rels")
        ]
        rel_targets = [name for name in names if name.startswith("word/") and name.endswith(".rels")]

        for name in xml_targets:
            try:
                xml_text = archive.read(name).decode("utf-8", errors="ignore")
            except Exception:
                continue

            xml_text = re.sub(r"</w:p>|</w:tr>|</w:tc>|</w:tbl>|</w:txbxContent>|</v:textbox>", "\n", xml_text)
            xml_text = re.sub(r"<[^>]+>", " ", xml_text)
            xml_text = re.sub(r"[ \t\r\f\v]+", " ", xml_text)
            xml_text = re.sub(r"\n+", "\n", xml_text)
            text_parts.extend(part.strip() for part in xml_text.split("\n") if part.strip())

        for name in rel_targets:
            try:
                rel_text = archive.read(name).decode("utf-8", errors="ignore")
            except Exception:
                continue
            for match in re.findall(r'Target="([^"]+)"', rel_text):
                if match.strip():
                    links.append(match.strip())

    return _dedupe_lines(text_parts), _dedupe_lines(links)


def _configure_tesseract_path() -> None:
    if pytesseract is None:
        return

    current_cmd = str(getattr(pytesseract.pytesseract, "tesseract_cmd", "") or "").strip()
    if current_cmd and os.path.exists(current_cmd):
        return

    for candidate in WINDOWS_TESSERACT_PATHS:
        if os.path.exists(candidate):
            pytesseract.pytesseract.tesseract_cmd = candidate
            return


def _ocr_score(text: str) -> tuple[int, int, int]:
    cleaned = _normalize_extracted_text(text)
    words = re.findall(r"\b[\w@.+-]{2,}\b", cleaned)
    emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", cleaned)
    lines = [line for line in cleaned.splitlines() if line.strip()]
    return (len(emails), len(words), len(lines))


def _ocr_image_variants(image: "Image.Image") -> list["Image.Image"]:
    if ImageOps is None or ImageFilter is None:
        return [image]

    grayscale = ImageOps.grayscale(image)
    autocontrast = ImageOps.autocontrast(grayscale)
    sharpened = autocontrast.filter(ImageFilter.SHARPEN)
    threshold = sharpened.point(lambda px: 255 if px > 170 else 0)
    return [image, autocontrast, sharpened, threshold]


def _ocr_pdf_bytes(data: bytes) -> str:
    if Image is None or pytesseract is None:
        return ""
    _configure_tesseract_path()

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception:
        return ""

    text_parts: list[str] = []
    for page in doc:
        try:
            pix = page.get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            best_text = ""
            best_score = (-1, -1, -1)
            for variant in _ocr_image_variants(image):
                for config in (
                    "--oem 3 --psm 6",
                    "--oem 3 --psm 4",
                    "--oem 3 --psm 11",
                ):
                    candidate = pytesseract.image_to_string(variant, config=config)
                    score = _ocr_score(candidate)
                    if score > best_score:
                        best_score = score
                        best_text = candidate
            ocr_text = best_text
        except Exception:
            ocr_text = ""

        if ocr_text.strip():
            text_parts.append(ocr_text)

    return _normalize_extracted_text("\n".join(text_parts))


def read_pdf_bytes(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    text = []
    links: list[str] = []
    for page in doc:
        text.append(page.get_text("text"))
        for link in page.get_links():
            uri = str(link.get("uri") or "").strip()
            if uri:
                links.append(uri)

    link_block = _format_links(links)
    extracted_text = _normalize_extracted_text("\n".join(text))

    if len(extracted_text) < PDF_OCR_FALLBACK_MIN_CHARS:
        ocr_text = _ocr_pdf_bytes(data)
        merged_text = "\n".join(part for part in [extracted_text, ocr_text] if part).strip()
    else:
        merged_text = extracted_text

    combined = "\n".join(part for part in [merged_text, link_block] if part).strip()
    return _normalize_extracted_text(combined)


def read_docx_bytes(data: bytes) -> str:
    bio = io.BytesIO(data)
    doc = Document(bio)
    parts = []
    links: list[str] = []
    for p in doc.paragraphs:
        parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = " ".join(paragraph.text for paragraph in cell.paragraphs if paragraph.text.strip())
                if cell_text.strip():
                    parts.append(cell_text)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

    for rel in doc.part.rels.values():
        if rel.reltype == RT.HYPERLINK:
            target = str(rel.target_ref or "").strip()
            if target:
                links.append(target)

    xml_parts, xml_links = _read_docx_xml_text(data)
    parts.extend(xml_parts)
    links.extend(xml_links)

    parts = _dedupe_lines(parts)
    link_block = _format_links(links)
    combined = "\n".join(part for part in ["\n".join(parts), link_block] if part).strip()
    return _normalize_extracted_text(combined)


def read_txt_bytes(data: bytes) -> str:
    return _normalize_extracted_text(data.decode("utf-8", errors="ignore"))


def load_uploaded_file(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return read_pdf_bytes(data)
    elif suffix == ".docx":
        return read_docx_bytes(data)
    elif suffix in [".txt", ".md"]:
        return read_txt_bytes(data)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
